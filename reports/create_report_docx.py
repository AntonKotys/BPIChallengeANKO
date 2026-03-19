"""Generate Word report from REPORT_MY_SECTIONS.md"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Title ──
p = doc.add_heading('Report Sections — Samat Tanikulov', level=0)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============================================================
# 1.4 Next Activities
# ============================================================
doc.add_heading('1.4 Next Activities', level=1)

doc.add_paragraph(
    'The next activity predictor is responsible for determining which branch a case '
    'follows at each XOR gateway. The implementation can be found in the '
    'task_1_4_next_activity.py file. Two versions were developed: a basic k-gram '
    'predictor with process model constraints, and an advanced approach that combines '
    'token replay with an enriched machine learning model.'
)

# ---- Basic Task ----
doc.add_heading('Basic Task', level=2)

doc.add_paragraph(
    'The basic predictor pairs a k-gram language model with a structural filter '
    'derived from the BPMN model. This means that the predictor can only output '
    'transitions that are valid according to the process model, which prevents '
    'impossible activity sequences from being generated during simulation.'
)

doc.add_paragraph('The training procedure works as follows:')
doc.add_paragraph(
    'We iterate over all cases in the filtered BPIC-17 event log and extract '
    'consecutive activity pairs.',
    style='List Number'
)
doc.add_paragraph(
    'For each pair (A, B), the predictor checks whether B is in the set of allowed '
    'successors of A according to the process model. If not, the pair is discarded.',
    style='List Number'
)
doc.add_paragraph(
    'Valid pairs are recorded under k-gram contexts of length 1 through k. '
    'We set k = 2 throughout.',
    style='List Number'
)
doc.add_paragraph(
    'The raw counts are converted to probability distributions using Laplace '
    'smoothing with pseudocount \u03b1 = 1.0:',
    style='List Number'
)

p = doc.add_paragraph()
run = p.add_run(
    'P(next | current, context) = (count(next) + \u03b1) / '
    '(total + \u03b1 \u00b7 |valid_successors|)'
)
run.italic = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    'Laplace smoothing was necessary because some branches at certain gateways '
    'appear rarely in the historical data. Without it, the predictor would assign '
    'zero probability to valid but infrequent paths.'
)

doc.add_paragraph(
    'At prediction time, the algorithm applies a fallback hierarchy. It first '
    'searches for the longest matching k-gram context in a per-activity table. If '
    'no match is found, it falls back to a global k-gram table, and finally to a '
    'flat distribution over all activities. Regardless of which level returns a '
    'result, the output is intersected with the set of valid successors and '
    'renormalized.'
)

doc.add_paragraph(
    'The major metric for this task is the average XOR difference between historical '
    'branching proportions and our predictor. Random branching produces a 37.5% '
    'difference, whilst our trained predictor achieves only 0.5%.'
)

# ---- Advanced Task ----
doc.add_heading('Advanced Task', level=2)

doc.add_paragraph(
    'For the advanced task, we developed two separate components: a token replay '
    'module to extract structurally correct decision point data, and a machine '
    'learning model that leverages both trace history and case-level attributes.'
)

p = doc.add_paragraph()
run = p.add_run('Token Replay')
run.bold = True

doc.add_paragraph(
    'A conformance-checking module was implemented in the task_1_4_token_replay.py '
    'file using the pm4py library. The module operates as follows:'
)
doc.add_paragraph(
    'The BPMN model is converted into a Petri net.',
    style='List Number'
)
doc.add_paragraph(
    'Every place with more than one outgoing arc is identified as a decision place. '
    'These correspond to XOR and OR gateways in the original model.',
    style='List Number'
)
doc.add_paragraph(
    'All 31,509 historical traces are replayed against the Petri net using '
    'pm4py\'s token replay algorithm.',
    style='List Number'
)
doc.add_paragraph(
    'Whenever a token arrives at a decision place and a transition fires, the '
    'module records the case identifier, the chosen transition, the full prefix of '
    'visible activities, and the set of available alternatives.',
    style='List Number'
)

doc.add_paragraph(
    'This procedure yielded 706,519 decision instances across 7 decision places. '
    'The data produced by token replay was considerably cleaner than what sequential '
    'pair counting produces, since the replay respects the concurrency and '
    'synchronization semantics of the Petri net rather than relying solely on the '
    'order of events in the flat log.'
)

doc.add_paragraph(
    'When the predictor is instantiated with use_token_replay=True, it first runs '
    'the basic k-gram fit as a fallback and then overwrites the transition counts '
    'with the replay-derived data. The prediction logic (context lookup, smoothing, '
    'sampling) remains the same.'
)

p = doc.add_paragraph()
run = p.add_run('Case-Attribute-Enriched ML Model')
run.bold = True

doc.add_paragraph(
    'The k-gram predictor, even when combined with token replay data, is still a '
    'frequency-based sampler. It cannot condition on case-specific attributes such as '
    'the requested loan amount or the applicant\'s credit score. In the context of a '
    'loan application process, these attributes are likely to influence whether an '
    'application gets approved, rejected, or sent for another round of offers.'
)

doc.add_paragraph(
    'To address this limitation, we trained a RandomForest classifier (500 trees, '
    'max depth 25) on an enriched feature set. In addition to the activity n-gram '
    'and temporal features (hour, weekday, elapsed seconds since case start), the '
    'following case-level attributes were included:'
)
features = [
    'RequestedAmount',
    'Encoded LoanGoal and ApplicationType',
    'Loop counter (number of times the current activity has appeared in the trace)',
    'Trace length up to the current point',
    'Cumulative count of offer-related events',
    'Most recent CreditScore and OfferedAmount observed in the trace',
    'Mean inter-event time',
    'Number of distinct resources that have worked on the case',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_paragraph(
    'This results in 11 case-level features in total. We initially trained a '
    'GradientBoosting model, which reached 63.2% accuracy on XOR decisions. '
    'Switching to RandomForest with the same feature set improved this to 81.4%. '
    'We attribute the difference to RandomForest\'s better handling of mixed '
    'categorical and continuous features and its lower tendency to overfit towards '
    'the majority class at imbalanced gateways.'
)

# ---- Evaluation ----
doc.add_heading('Evaluation', level=2)

doc.add_paragraph(
    'One complication arose during evaluation. Our BPMN model uses activity '
    'clustering, meaning that certain raw activities are merged into single nodes '
    '(e.g., "W_Complete application" and "A_Concept" become "W_Complete application '
    '& A_Concept"). The event log, however, retains the original unclustered names. '
    'As a result, 5 out of 7 XOR gateways reference activity names that do not '
    'appear as consecutive transitions in the raw log. On the raw data, we could '
    'therefore only evaluate 2 gateways. This was a team-wide modelling decision and '
    'could not be changed without affecting other parts of the simulation.'
)

doc.add_paragraph(
    'To obtain a more complete evaluation, we ran an additional experiment in which '
    'the activity mapping was applied to the event log before training. Raw '
    'activities were replaced by their clustered equivalents, and consecutive '
    'duplicates within each trace were collapsed. This brought all 7 gateways into '
    'scope and increased the number of evaluable XOR decisions from approximately '
    '8,200 to over 18,400.'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Table X: Results on mapped event log (all 7 gateways)')
run.bold = True

table = doc.add_table(rows=4, cols=4, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Metric', 'Random', 'Basic (k-gram)', 'Best (RF enriched)']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for paragraph in table.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

data = [
    ['Avg XOR diff vs historical', '37.5%', '0.5%', 'n/a'],
    ['XOR prediction accuracy', '~33-50%', '51.1%', '81.4%'],
    ['Gateways evaluated', '7/7', '7/7', '7/7'],
]
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        table.rows[r+1].cells[c].text = val

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Table X: Per-gateway breakdown')
run.bold = True

table2 = doc.add_table(rows=5, cols=4, style='Light Shading Accent 1')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
headers2 = ['Gateway', 'Basic', 'RF enriched', 'Decisions']
for i, h in enumerate(headers2):
    table2.rows[0].cells[i].text = h
    for paragraph in table2.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

gw_data = [
    ['A_Submitted', '67.3%', '100.0%', '4,136'],
    ['O_Sent (mail and online)', '59.5%', '86.0%', '858'],
    ['A_Create Application', '54.4%', '83.9%', '6,302'],
    ['W_Call after offers & A_Complete', '37.3%', '69.6%', '7,090'],
]
for r, row_data in enumerate(gw_data):
    for c, val in enumerate(row_data):
        table2.rows[r+1].cells[c].text = val

doc.add_paragraph()
doc.add_paragraph(
    'The basic k-gram predictor achieves a very low distribution error (0.5%), '
    'meaning that the simulated process visits each branch in approximately the '
    'same proportions as the real process. For a simulation whose primary goal is '
    'to reproduce aggregate behavior, this is the more relevant metric.'
)

doc.add_paragraph(
    'For per-instance prediction accuracy, the enriched RandomForest significantly '
    'outperforms the k-gram baseline (81.4% vs 51.1%). The most difficult gateway '
    'is the three-way split at "W_Call after offers & A_Complete," where the '
    'historical distribution is approximately 62/26/13 between validate, new offer, '
    'and cancel. The k-gram predictor barely outperforms random selection at this '
    'gateway (37.3%), whilst the RandomForest reaches 69.6%.'
)

doc.add_paragraph(
    'We also experimented with different window sizes (k = 3, 5, and 7). The '
    'accuracy differences were negligible (81.0% to 81.4%), which suggests that '
    'the case-level attributes carry more predictive signal than longer activity '
    'histories.'
)

# ============================================================
# 2.1 Resource Allocation
# ============================================================
doc.add_heading('2.1 Resource Allocation — Advanced Methods', level=1)

doc.add_paragraph(
    'In addition to the basic heuristics (random, round-robin, earliest-available), '
    'we implemented two advanced strategies.'
)

doc.add_heading('K-Batching', level=2)
doc.add_paragraph(
    'Based on Zeng and Zhao (2005) [3]. Instead of immediate assignment, the allocator '
    'accumulates K tasks (default K=5), then solves a Parallel Machines Scheduling '
    'Problem using either:'
)
doc.add_paragraph(
    'LPT heuristic: Tasks sorted by decreasing duration, greedily assigned to '
    'earliest-free resource (4/3-approximation for makespan).',
    style='List Bullet'
)
doc.add_paragraph(
    'Hungarian algorithm: Optimal one-to-one assignment via '
    'scipy.optimize.linear_sum_assignment (O(n³)).',
    style='List Bullet'
)
doc.add_paragraph(
    'A timeout (default 3600s) prevents starvation. Implemented in '
    'ResourceAllocator/BatchAllocator.py, wrapping the base ResourceAllocatorAlgo.'
)

doc.add_heading('SVFA (Score-based Value Function Approximation)', level=2)
doc.add_paragraph(
    'Based on Middelhuis et al. (2025) [2]. Scores every (resource, task) pair '
    'using a learned linear combination of six features:'
)

p = doc.add_paragraph()
run = p.add_run(
    'Score(r, k) = w₁·MeanAssignment + w₂·VarAssignment + w₃·ActivityRank '
    '+ w₄·ResourceRank − w₅·ProbFin − w₆·QueueLength'
)
run.italic = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    'If no pair scores below threshold w₇, assignments are postponed. Weights are '
    'trained via Bayesian optimization (scikit-optimize, gp_minimize) minimizing '
    'mean cycle time. Implemented in ResourceAllocator/SVFAllocator.py.'
)
doc.add_paragraph(
    'Both strategies are integrated into the simulation engine as '
    'allocation_strategy="k_batch" / "svfa".'
)

# ============================================================
# 2.2 Evaluation
# ============================================================
doc.add_heading('2.2 Evaluation — Advanced Allocation Strategies', level=1)

doc.add_paragraph('Scaling experiment: 100, 500, and 1,000 cases, all five strategies.')

p = doc.add_paragraph()
run = p.add_run('Results at 1,000 Cases:')
run.bold = True

table3 = doc.add_table(rows=6, cols=4, style='Light Shading Accent 1')
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
headers3 = ['Strategy', 'Cycle Time (days)', 'Activity Delay (h)', 'Fairness (Jain)']
for i, h in enumerate(headers3):
    table3.rows[0].cells[i].text = h
    for paragraph in table3.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

eval_data = [
    ['Earliest Available', '19.45', '8.98', '0.60'],
    ['Round Robin', '26.37', '20.50', '0.61'],
    ['Random', '29.69', '23.31', '0.58'],
    ['K-Batch', '58.77', '99.66', '0.91'],
    ['SVFA', '86.06', '173.24', '0.88'],
]
for r, row_data in enumerate(eval_data):
    for c, val in enumerate(row_data):
        table3.rows[r+1].cells[c].text = val

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Key finding: efficiency-fairness trade-off. ')
run.bold = True
p.add_run(
    'K-Batch and SVFA achieve ~50% higher fairness (Jain index 0.88–0.91 vs 0.60) '
    'but at the cost of longer cycle times. At low load (100 cases), SVFA achieves '
    'the lowest cycle time (13.61 days vs 14.30 for earliest-available), demonstrating '
    'that intelligent assignment provides genuine value when coordination overhead is '
    'negligible. Under high load, batching/scoring delays compound in a cascading effect.'
)

p = doc.add_paragraph()
run = p.add_run('Practical recommendation: ')
run.bold = True
p.add_run(
    'Use earliest-available for SLA-critical processes; use K-Batch/SVFA when workload '
    'fairness matters (preventing burnout). A hybrid approach switching strategies '
    'based on load could combine both benefits.'
)

# ============================================================
# References
# ============================================================
doc.add_heading('References', level=1)
doc.add_paragraph(
    '[2] J. Middelhuis et al., "Learning policies for resource allocation in '
    'business processes," Inf. Syst., vol. 128, art. 102492, 2025.'
)
doc.add_paragraph(
    '[3] D. D. Zeng and J. L. Zhao, "Effective role resolution in workflow '
    'management," INFORMS J. Comput., vol. 17, no. 3, pp. 374–387, 2005.'
)

# Save
output_path = '/Users/s.tanikulov/PycharmProjects/BPIChallengeANKO/Report_Samat_Tanikulov.docx'
doc.save(output_path)
print(f"Saved to {output_path}")
