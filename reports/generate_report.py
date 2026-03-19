"""
Generate Word report for Task 1.4: Next Activity Prediction
Analysis of Basic and Advanced implementation, comparison with other teams.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os


def set_cell_shading(cell, color_hex):
    from docx.oxml.ns import qn
    from lxml import etree
    shading = etree.SubElement(cell._tc.get_or_add_tcPr(), qn('w:shd'))
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    return table


def generate_report():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # ================================================================
    # TITLE PAGE
    # ================================================================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Task 1.4: Next Activity Prediction\n')
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Analysis Report — Basic & Advanced Implementation\n')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x4F, 0x81, 0xBD)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        'Business Process Simulation, Prediction and Optimization\n'
        'BPI Challenge 2017 Dataset\n\n'
        'Technical University of Munich\n'
        'Chair of Information Systems and Business Process Management\n'
    )
    run.font.size = Pt(11)

    doc.add_page_break()

    # ================================================================
    # TABLE OF CONTENTS (placeholder)
    # ================================================================
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Introduction',
        '2. Assignment Requirements (Task 1.4)',
        '3. Implementation Analysis',
        '   3.1 Basic Task: Branching Probabilities (k-gram Model)',
        '   3.2 Advanced Task: ML-based Prediction (GradientBoosting)',
        '   3.3 Integration with Simulation Engine',
        '   3.4 Evaluation & Metrics',
        '4. Compliance Assessment',
        '5. Comparison with Other Teams',
        '   5.1 Team 1: C4.5 Decision Tree',
        '   5.2 Team 2: Bigram Model',
        '   5.3 Team 3: Random Forest',
        '   5.4 Team 4: MLP Classifier',
        '   5.5 Comparison Summary Table',
        '6. Strengths & Weaknesses of Current Implementation',
        '7. Recommendations for Improvement',
        '8. Conclusion',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)

    doc.add_page_break()

    # ================================================================
    # 1. INTRODUCTION
    # ================================================================
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'This report provides a detailed analysis of the implementation of Task 1.4 '
        '(Next Activity Prediction) within the Business Process Simulation project based '
        'on the BPI Challenge 2017 dataset. The task requires predicting which branch to '
        'take at XOR gateways in a BPMN process model during discrete-event simulation.'
    )
    doc.add_paragraph(
        'The implementation is contained primarily in the following files:'
    )

    files_table = [
        ['task_1_4_next_activity.py', '~781 lines', 'Core predictor: ExpertActivityPredictor class with Basic (k-gram) and Advanced (GradientBoosting) modes'],
        ['simulation_engine_core_V1.5.py', '~797 lines', 'Latest simulation engine integrating the predictor into route_next() for XOR/OR gateway decisions'],
        ['task_1_4_metrics.py', '~390 lines', 'Evaluation framework: transition validity, XOR branching comparison (sim vs. historical)'],
        ['test_task_1_4.py', '~353 lines', 'Unit test suite (6 tests): fit/predict, XOR restriction, sampling, context fallback, integration, serialization'],
        ['verify_task_1_4.py', '~292 lines', 'Verification: trace-history influence, enabled_next constraint, simulation output validation'],
        ['branch_comparison_report.py', '~293 lines', 'Gateway comparison report generator (historical vs. trained vs. random)'],
    ]

    add_styled_table(doc,
        ['File', 'Size', 'Description'],
        files_table
    )
    doc.add_paragraph()

    # ================================================================
    # 2. ASSIGNMENT REQUIREMENTS
    # ================================================================
    doc.add_heading('2. Assignment Requirements (Task 1.4)', level=1)

    doc.add_heading('2.1 Basic Requirement', level=2)
    doc.add_paragraph(
        'As stated in the assignment: "To predict at XOR gateways which branch to take, '
        'you must at least implement a branching decision based on branching probabilities '
        'learned from the event log (the challenge here lies in identifying which traces are '
        'relevant for which branches). As a basic approach, this can be based on looking only '
        'at the preceding activity/activities and the succeeding activities of a decision point."'
    )

    doc.add_heading('2.2 Advanced Requirement', level=2)
    doc.add_paragraph(
        'As stated in the assignment: "Identify the data for the decision points via token replay. '
        'Train a predictive model that takes, e.g., the history of a trace into account."'
    )

    doc.add_page_break()

    # ================================================================
    # 3. IMPLEMENTATION ANALYSIS
    # ================================================================
    doc.add_heading('3. Implementation Analysis', level=1)

    # 3.1 BASIC
    doc.add_heading('3.1 Basic Task: Branching Probabilities (k-gram Model)', level=2)

    doc.add_heading('3.1.1 Approach', level=3)
    doc.add_paragraph(
        'The Basic implementation uses a process-model-aware k-gram model. The core class is '
        'ExpertActivityPredictor (mode="basic"). Key design decisions:'
    )

    basic_points = [
        ('K-gram Context', 'Learns transition probabilities P(next | current_activity, preceding_k_activities). '
         'The parameter basic_context_k (default=2) controls how many preceding activities form the context.'),
        ('Process Model Awareness', 'During training, only transitions that are VALID according to the process model '
         'are counted. If a transition current_act → next_act is not in PROCESS_MODEL[current_act], it is skipped. '
         'This is a key differentiator from naive frequency counting.'),
        ('Laplace Smoothing', 'All probability computations use Laplace smoothing (alpha=1.0 by default) to handle '
         'unseen transitions gracefully. This ensures no valid branch gets zero probability.'),
        ('Hierarchical Fallback', 'If the full k-context is not found, the model falls back to shorter contexts '
         '(k-1, k-2, ..., 1), and ultimately to a global distribution. This ensures predictions are always available.'),
        ('Per-Activity Tables', 'Transition counts are organized per current_activity, then per context length k, '
         'then per context tuple. This enables efficient lookup: transition_counts[current_act][k][context][next_act].'),
    ]

    for title_text, desc in basic_points:
        p = doc.add_paragraph()
        run = p.add_run(f'{title_text}: ')
        run.bold = True
        p.add_run(desc)

    doc.add_heading('3.1.2 Training Process (_fit_basic_kgram_process_aware)', level=3)
    doc.add_paragraph(
        'The training iterates over all cases in the event log. For each consecutive pair '
        '(activities[i], activities[i+1]):'
    )
    steps = [
        'Check if the transition is valid per the process model',
        'If valid, for each context length k = 1..basic_context_k, extract the k preceding activities as context',
        'Increment transition_counts[current_act][k][context][next_act]',
        'Convert counts to smoothed probabilities using _compute_smoothed_probs()',
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Number')

    doc.add_heading('3.1.3 Prediction Process', level=3)
    doc.add_paragraph(
        'At prediction time (get_next_activity_distribution or sample_next_activity):'
    )
    pred_steps = [
        'Determine current_activity (last in prefix)',
        'Look up per-activity transition table for current_activity',
        'Try longest context first (k=basic_context_k), fall back to shorter',
        'Filter to enabled_next (XOR outgoing arcs) and re-normalize with smoothing',
        'Sample from the resulting distribution',
    ]
    for s in pred_steps:
        doc.add_paragraph(s, style='List Number')

    doc.add_heading('3.1.4 Assessment of Basic Task', level=3)
    p = doc.add_paragraph()
    run = p.add_run('Verdict: FULLY IMPLEMENTED. ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    p.add_run(
        'The basic task requires "branching probabilities learned from the event log" based on '
        '"preceding activity/activities and succeeding activities of a decision point." '
        'The k-gram model does exactly this: it learns P(next | current, context) from historical data '
        'and applies it at XOR gateways. The process model awareness further ensures only valid '
        'branches are considered, which exceeds the basic requirement.'
    )

    doc.add_paragraph()

    # 3.2 ADVANCED
    doc.add_heading('3.2 Advanced Task: ML-based Prediction (GradientBoosting)', level=2)

    doc.add_heading('3.2.1 Approach', level=3)
    doc.add_paragraph(
        'The Advanced implementation uses a GradientBoostingClassifier from scikit-learn, trained on '
        'contextual features extracted from the event log. When mode="advanced", the _fit_advanced_ml() '
        'method is called during training.'
    )

    adv_points = [
        ('Feature Engineering', 'For each transition in the log, the following features are extracted:\n'
         '  - N-gram history: the last window_size (default=5) encoded activity indices\n'
         '  - Hour of day (0-23)\n'
         '  - Day of week (0-6)\n'
         '  - Elapsed time since case start (seconds)\n'
         'Total feature vector: window_size + 3 dimensions.'),
        ('Activity Encoding', 'Activities are encoded using sklearn LabelEncoder. Unknown activities get index -1. '
         'Shorter histories are padded with -1.'),
        ('Model', 'GradientBoostingClassifier(n_estimators=50, max_depth=3, random_state=42). '
         'A relatively lightweight model to avoid overfitting.'),
        ('Train/Validation Split', 'GroupShuffleSplit with test_size=0.2, grouped by case_id. '
         'This ensures entire cases go to either train or validation (no data leakage).'),
        ('Prediction with Fallback', 'At prediction time, the ML model predicts the most likely next activity. '
         'If the predicted activity is not in enabled_next (invalid per process model), '
         'it falls back to the basic k-gram sampling. This ensures process model compliance.'),
    ]

    for title_text, desc in adv_points:
        p = doc.add_paragraph()
        run = p.add_run(f'{title_text}: ')
        run.bold = True
        p.add_run(desc)

    doc.add_heading('3.2.2 Assessment of Advanced Task', level=3)
    p = doc.add_paragraph()
    run = p.add_run('Verdict: IMPLEMENTED, but with caveats. ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)

    doc.add_paragraph(
        'The advanced task requires: "Identify the data for the decision points via token replay. '
        'Train a predictive model that takes, e.g., the history of a trace into account."'
    )
    doc.add_paragraph('What IS implemented:')
    implemented_items = [
        'A GradientBoosting ML model that uses trace history (n-gram of activities) and contextual features (time, weekday, elapsed)',
        'Proper train/validation split with GroupShuffleSplit (no data leakage)',
        'Fallback to basic model when ML prediction is invalid per process model',
        'Validation accuracy is printed during training',
    ]
    for item in implemented_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph('What is MISSING or could be improved:')
    missing_items = [
        'Token replay is NOT explicitly implemented. The assignment asks to "identify the data for the decision points '
        'via token replay." The current implementation extracts transitions from sequential event log directly, '
        'not via conformance-checking-based token replay. This is a significant gap.',
        'The advanced mode is NOT used in the latest simulation_engine_core_V1.5.py (mode="basic" in __main__). '
        'It exists as code but has not been activated in the final simulation runs.',
        'The ML model predicts the single most likely class (argmax), not a probability distribution. '
        'For stochastic simulation, probability-based sampling would be more appropriate.',
        'No hyperparameter tuning or model selection comparison is provided.',
        'Feature engineering is minimal — no case-level attributes (e.g., loan amount, applicant type) are used, '
        'despite being available in BPI 2017.',
    ]
    for item in missing_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # 3.3 INTEGRATION
    doc.add_heading('3.3 Integration with Simulation Engine', level=2)

    doc.add_paragraph(
        'The predictor is integrated into the SimulationEngine via the route_next() method '
        '(simulation_engine_core_V1.5.py, lines 233-344). The integration handles three cases:'
    )

    integration_items = [
        ('XOR Gateways (lines 324-336)',
         'Uses predictor.sample_next_activity(prefix_activities=trace, enabled_next=outgoing) to '
         'sample one activity from the learned distribution. Falls back to random.choice() if predictor '
         'returns None.'),
        ('OR Gateway: O_Sent (mail and online) (lines 262-283)',
         'Special handling with offer count limit (MAX_OFFERS_PER_CASE=3). Uses inclusive OR among '
         'remaining branches. Does NOT use the predictor — uses random sampling.'),
        ('OR Gateway: W_Call after offers & A_Complete (lines 287-322)',
         'Special handling where cancellation is exclusive (if cancel chosen, only cancel). Uses '
         'predictor.get_next_activity_distribution() to estimate cancel probability (~4.7% default). '
         'Non-cancel branches use inclusive OR random sampling.'),
    ]

    for title_text, desc in integration_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{title_text}: ')
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph(
        '\nThe case trace history is tracked in self.case_traces (Dict[str, List[str]]), '
        'updated in log_event(). This provides the prefix for the predictor.'
    )

    # 3.4 EVALUATION
    doc.add_heading('3.4 Evaluation & Metrics', level=2)
    doc.add_paragraph(
        'A comprehensive evaluation framework was built (task_1_4_metrics.py and branch_comparison_report.py):'
    )

    eval_items = [
        'Transition validity checking against the process model',
        'XOR branching probability comparison (simulated vs. historical)',
        'Gateway comparison report: Historical vs. Trained vs. Random',
    ]
    for item in eval_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        '\nFrom the gateway_comparison_report.txt (generated with 100 simulated cases vs. 31,509 historical):'
    )

    eval_results = [
        ['A_Create Application', 'XOR', '0.0%', '106.0%', 'Trained perfectly matches (100% → A_Submitted)'],
        ['A_Submitted', 'XOR', '0.0%', '85.1%', 'Trained perfectly matches (100% → W_Handle leads)'],
        ['O_Returned', 'XOR', '100.0%', '100.0%', 'Neither matches — historical mapping issue'],
        ['A_Incomplete', 'XOR', '100.0%', '100.0%', 'Neither matches — historical mapping issue'],
        ['A_Validating', 'XOR', '100.0%', '100.0%', 'Neither matches — historical mapping issue'],
        ['W_Call after offers...', 'OR', '100.0%', '100.0%', 'Neither matches — historical mapping issue'],
        ['TOTAL', '', '400.0%', '591.1%', 'Trained 32.3% closer to historical than random'],
    ]

    add_styled_table(doc,
        ['Gateway', 'Type', 'Trained Dev.', 'Random Dev.', 'Notes'],
        eval_results
    )

    doc.add_paragraph(
        '\nNote: The high deviation for O_Returned, A_Incomplete, A_Validating, and the OR gateway '
        'is caused by activity name mapping differences between the historical log and the simulation model '
        '(e.g., "W_Call incomplete files" in historical maps to "W_Validate application & A_Validating" in simulation). '
        'For the two gateways where mapping is correct (A_Create Application, A_Submitted), the trained predictor '
        'achieves 0% deviation — a perfect match.'
    )

    doc.add_page_break()

    # ================================================================
    # 4. COMPLIANCE ASSESSMENT
    # ================================================================
    doc.add_heading('4. Compliance Assessment', level=1)

    compliance = [
        ['Basic: Learn branching probabilities from event log', 'YES', 'K-gram model learns P(next|current, context) from historical transitions'],
        ['Basic: Based on preceding activities', 'YES', 'Uses k preceding activities as context (configurable k)'],
        ['Basic: Identify relevant traces for branches', 'YES', 'Process-model-aware filtering: only valid transitions counted'],
        ['Advanced: Token replay for decision points', 'PARTIAL', 'Decision points identified via process model, but no explicit token replay (conformance checking)'],
        ['Advanced: Predictive model with trace history', 'YES', 'GradientBoosting with n-gram history + contextual features'],
        ['Integration with simulation engine', 'YES', 'Fully integrated into route_next() for XOR and OR gateways'],
        ['Evaluation/comparison', 'YES', 'Comprehensive metrics: validity, XOR comparison, gateway report'],
        ['Test suite', 'YES', '6 unit tests + 3 verification tests'],
    ]

    add_styled_table(doc,
        ['Requirement', 'Met?', 'Evidence'],
        compliance
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Overall Assessment: ')
    run.bold = True
    p.add_run(
        'The Basic task is fully and correctly implemented. The Advanced task is implemented in code '
        'but has two gaps: (1) token replay is not explicitly used, and (2) the advanced mode is not '
        'activated in the final simulation run (only basic mode is used in __main__). '
        'The implementation quality is high, with good software engineering practices (modular design, '
        'tests, evaluation framework, serialization).'
    )

    doc.add_page_break()

    # ================================================================
    # 5. COMPARISON WITH OTHER TEAMS
    # ================================================================
    doc.add_heading('5. Comparison with Other Teams', level=1)
    doc.add_paragraph(
        'Four other teams implemented Task 1.4 using different approaches. Below is a detailed '
        'analysis of each, including advantages and disadvantages.'
    )

    # Team 1
    doc.add_heading('5.1 Team 1: C4.5 Decision Tree', level=2)
    doc.add_paragraph(
        'The C4.5 algorithm (also known as J48 in Weka, implemented as DecisionTreeClassifier with '
        'entropy criterion in scikit-learn) builds a decision tree that splits on the most informative features.'
    )
    p = doc.add_paragraph()
    run = p.add_run('Approach: ')
    run.bold = True
    p.add_run(
        'Train a C4.5 decision tree on activity sequences to predict the next activity at XOR gateways. '
        'Features likely include encoded activity history and possibly case attributes.'
    )

    doc.add_paragraph('Advantages:', style='List Bullet')
    adv_c45 = [
        'Highly interpretable — the decision tree can be visualized and explained to stakeholders',
        'Handles both categorical and numerical features natively',
        'Information gain (entropy) criterion is well-suited for classification tasks with multiple classes',
        'Fast inference (tree traversal is O(depth))',
        'Can reveal which features are most important for branching decisions',
    ]
    for a in adv_c45:
        doc.add_paragraph(a, style='List Bullet 2')

    doc.add_paragraph('Disadvantages:', style='List Bullet')
    dis_c45 = [
        'Prone to overfitting, especially with deep trees and noisy process data',
        'Deterministic predictions — always predicts the same class for the same input (no stochastic sampling)',
        'Cannot easily capture sequential dependencies (history as features must be explicitly engineered)',
        'Sensitive to class imbalance (common in process mining where some branches are rare)',
        'C4.5 specifically handles missing values but adds complexity',
    ]
    for d in dis_c45:
        doc.add_paragraph(d, style='List Bullet 2')

    # Team 2
    doc.add_heading('5.2 Team 2: Bigram Model', level=2)
    doc.add_paragraph(
        'A bigram model uses pairs of consecutive activities (A, B) to predict the next activity after B.'
    )
    p = doc.add_paragraph()
    run = p.add_run('Approach: ')
    run.bold = True
    p.add_run(
        'Count pairs of consecutive activities in the event log, compute P(next | current) as '
        'frequency ratios. At XOR gateways, sample from these bigram probabilities.'
    )

    doc.add_paragraph('Advantages:', style='List Bullet')
    adv_bigram = [
        'Extremely simple to implement and understand',
        'Computationally efficient — O(n) training, O(1) prediction',
        'Naturally produces probability distributions (good for stochastic simulation)',
        'No hyperparameters to tune',
        'Directly models what the basic task requires',
    ]
    for a in adv_bigram:
        doc.add_paragraph(a, style='List Bullet 2')

    doc.add_paragraph('Disadvantages:', style='List Bullet')
    dis_bigram = [
        'Only considers one preceding activity — ignores longer history',
        'Cannot capture complex patterns (e.g., "if activity X happened 3 steps ago, branch differently")',
        'No contextual features (time, resources, case attributes)',
        'Suffers from data sparsity for rare transitions',
        'Essentially equivalent to k=1 of the k-gram model in this implementation',
    ]
    for d in dis_bigram:
        doc.add_paragraph(d, style='List Bullet 2')

    # Team 3
    doc.add_heading('5.3 Team 3: Random Forest', level=2)
    doc.add_paragraph(
        'Random Forest is an ensemble of decision trees, each trained on a bootstrap sample '
        'of the data with random feature subsets.'
    )
    p = doc.add_paragraph()
    run = p.add_run('Approach: ')
    run.bold = True
    p.add_run(
        'Train a Random Forest classifier with features derived from activity sequences and '
        'possibly contextual information to predict the next activity.'
    )

    doc.add_paragraph('Advantages:', style='List Bullet')
    adv_rf = [
        'More robust than single decision tree — reduces overfitting via bagging',
        'Provides class probability estimates (good for stochastic sampling)',
        'Handles high-dimensional feature spaces well',
        'Feature importance ranking helps understand branching factors',
        'Relatively fast training and prediction',
        'Robust to noise and outliers in process data',
    ]
    for a in adv_rf:
        doc.add_paragraph(a, style='List Bullet 2')

    doc.add_paragraph('Disadvantages:', style='List Bullet')
    dis_rf = [
        'Less interpretable than single decision tree',
        'Can be computationally expensive with many trees',
        'May overfit with too many trees on small event logs',
        'Feature engineering for sequential data is still required (like for any ML model)',
        'Memory-intensive for large forests',
    ]
    for d in dis_rf:
        doc.add_paragraph(d, style='List Bullet 2')

    # Team 4
    doc.add_heading('5.4 Team 4: MLP Classifier (Neural Network)', level=2)
    doc.add_paragraph(
        'Multi-Layer Perceptron (MLP) is a feedforward neural network with one or more hidden layers.'
    )
    p = doc.add_paragraph()
    run = p.add_run('Approach: ')
    run.bold = True
    p.add_run(
        'Train an MLP classifier on encoded activity features to predict the next branch at '
        'XOR gateways. The network learns non-linear decision boundaries.'
    )

    doc.add_paragraph('Advantages:', style='List Bullet')
    adv_mlp = [
        'Can learn complex non-linear patterns in branching decisions',
        'Naturally outputs class probabilities via softmax (good for stochastic simulation)',
        'Flexible architecture — can add layers for more complex patterns',
        'Can incorporate diverse feature types (activity encodings, numerical attributes, etc.)',
        'With proper architecture, can capture sequential dependencies',
    ]
    for a in adv_mlp:
        doc.add_paragraph(a, style='List Bullet 2')

    doc.add_paragraph('Disadvantages:', style='List Bullet')
    dis_mlp = [
        'Black-box model — no interpretability',
        'Requires careful hyperparameter tuning (learning rate, layers, neurons, regularization)',
        'Sensitive to feature scaling and preprocessing',
        'Can overfit on small datasets (BPI 2017 has ~31K cases — moderate size)',
        'Slower training than tree-based methods',
        'Not inherently designed for sequential data (unlike LSTM/GRU)',
    ]
    for d in dis_mlp:
        doc.add_paragraph(d, style='List Bullet 2')

    # 5.5 COMPARISON TABLE
    doc.add_heading('5.5 Comparison Summary Table', level=2)

    comp_headers = ['Criterion', 'Your Impl. (k-gram + GB)', 'C4.5 Tree', 'Bigram', 'Random Forest', 'MLP']
    comp_rows = [
        ['Approach Type', 'Statistical + ML', 'Decision Tree', 'Statistical', 'Ensemble (Trees)', 'Neural Network'],
        ['Complexity', 'Medium', 'Medium', 'Low', 'Medium-High', 'High'],
        ['Interpretability', 'High (basic) / Low (adv.)', 'High', 'Very High', 'Medium', 'Low'],
        ['Probability Output', 'Yes (natural)', 'Leaf frequency', 'Yes (natural)', 'Yes (voting)', 'Yes (softmax)'],
        ['History Awareness', 'Yes (k-gram)', 'Via features', 'No (k=1)', 'Via features', 'Via features'],
        ['Process Model Aware', 'Yes (key feature)', 'Unlikely', 'Unlikely', 'Unlikely', 'Unlikely'],
        ['Contextual Features', 'Yes (advanced)', 'Yes', 'No', 'Yes', 'Yes'],
        ['Training Speed', 'Fast', 'Fast', 'Very Fast', 'Moderate', 'Slow'],
        ['Overfitting Risk', 'Low (basic) / Med (adv.)', 'High', 'Low', 'Low-Medium', 'High'],
        ['Stochastic Sampling', 'Yes', 'Need adaptation', 'Yes', 'Yes', 'Yes'],
        ['Handles Rare Branches', 'Yes (Laplace smooth)', 'Poor', 'Poor', 'Moderate', 'Poor'],
    ]

    add_styled_table(doc, comp_headers, comp_rows)

    doc.add_page_break()

    # ================================================================
    # 6. STRENGTHS & WEAKNESSES
    # ================================================================
    doc.add_heading('6. Strengths & Weaknesses of Current Implementation', level=1)

    doc.add_heading('6.1 Strengths', level=2)
    strengths = [
        ('Process Model Awareness', 'Unlike any other team\'s approach, the implementation '
         'constrains predictions to valid process model transitions during both training and prediction. '
         'This is a significant design advantage that prevents impossible transitions.'),
        ('Hierarchical Fallback', 'The k-gram model gracefully degrades from full context to shorter context '
         'to global distribution. This ensures predictions are always available, even for unseen contexts.'),
        ('Dual Mode Design', 'Having both basic (k-gram) and advanced (GradientBoosting) in a single class '
         'with clean API design is good software engineering.'),
        ('Comprehensive Testing', '6 unit tests + 3 verification tests + evaluation framework + comparison report. '
         'This level of testing is unusual for an academic project and shows maturity.'),
        ('OR Gateway Handling', 'Special handling for OR gateways (offer limits, exclusive cancellation) '
         'shows deep understanding of the process semantics.'),
        ('Laplace Smoothing', 'Proper handling of unseen transitions prevents zero-probability issues.'),
        ('Serialization', 'JSON save/load capability enables model persistence and reproducibility.'),
    ]
    for title_text, desc in strengths:
        p = doc.add_paragraph()
        run = p.add_run(f'{title_text}: ')
        run.bold = True
        p.add_run(desc)

    doc.add_heading('6.2 Weaknesses', level=2)
    weaknesses = [
        ('No Token Replay', 'The advanced requirement explicitly asks for token replay. The implementation '
         'identifies decision points from the process model directly, not via conformance-checking-based token replay. '
         'Using pm4py\'s token-based replay would address this gap.'),
        ('Advanced Mode Not Used', 'The GradientBoosting model is implemented but never used in the actual simulation '
         '(simulation_engine_core_V1.5.py uses mode="basic"). This weakens the advanced task claim.'),
        ('ML Model Does Not Return Distribution', 'The advanced model uses model.predict() (argmax) instead of '
         'model.predict_proba(). For stochastic simulation, probability-based sampling would be more faithful.'),
        ('Activity Mapping Issues', 'The gateway comparison report shows 100% deviation for 4 out of 6 gateways, '
         'caused by activity name mapping issues between historical and simulated data. This significantly '
         'undermines the evaluation results.'),
        ('No Case Attributes', 'BPI 2017 contains rich case attributes (loan amount, application type, etc.) '
         'that could improve prediction accuracy. Neither mode uses them.'),
        ('Hardcoded Process Logic', 'The OR gateway handling in route_next() is hardcoded for specific activities '
         '("W_Call after offers & A_Complete"), not process-model-agnostic.'),
    ]
    for title_text, desc in weaknesses:
        p = doc.add_paragraph()
        run = p.add_run(f'{title_text}: ')
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ================================================================
    # 7. RECOMMENDATIONS
    # ================================================================
    doc.add_heading('7. Recommendations for Improvement', level=1)

    doc.add_paragraph(
        'Based on the analysis, the following improvements are recommended, ordered by priority:'
    )

    recs = [
        ('HIGH PRIORITY: Activate Advanced Mode',
         'Change the simulation __main__ to use mode="advanced" for at least one simulation run. '
         'Compare basic vs. advanced results quantitatively. This would validate the advanced implementation.'),
        ('HIGH PRIORITY: Add Token Replay',
         'Use pm4py\'s token_replay module to identify decision points from the event log. '
         'This directly addresses the assignment\'s advanced requirement: '
         '"Identify the data for the decision points via token replay." '
         'Example: pm4py.conformance_diagnostics_token_based_replay(log, net, im, fm).'),
        ('MEDIUM PRIORITY: Use predict_proba() in Advanced Mode',
         'Replace self.model.predict(features) with self.model.predict_proba(features) to get '
         'a probability distribution over next activities. Sample from this distribution for '
         'stochastic simulation behavior.'),
        ('MEDIUM PRIORITY: Fix Activity Mapping',
         'Ensure the historical data mapping in task_1_4_metrics.py perfectly aligns with the simulation model. '
         'The current 100% deviation for most gateways makes evaluation unreliable.'),
        ('LOW PRIORITY: Add Case Attributes',
         'Incorporate BPI 2017 case attributes (RequestedAmount, ApplicationType, etc.) as features '
         'in the advanced model. These could significantly improve prediction accuracy for XOR decisions.'),
        ('LOW PRIORITY: Hyperparameter Tuning',
         'Run a grid search or Bayesian optimization over GradientBoosting hyperparameters '
         '(n_estimators, max_depth, learning_rate) with cross-validation.'),
    ]

    for title_text, desc in recs:
        p = doc.add_paragraph()
        run = p.add_run(f'{title_text}\n')
        run.bold = True
        p.add_run(desc)
        p.paragraph_format.space_after = Pt(8)

    doc.add_page_break()

    # ================================================================
    # 8. CONCLUSION
    # ================================================================
    doc.add_heading('8. Conclusion', level=1)

    doc.add_paragraph(
        'The implementation of Task 1.4 demonstrates a solid understanding of next activity prediction '
        'in business process simulation. The key findings are:'
    )

    conclusions = [
        'The Basic task is fully and correctly implemented with a process-model-aware k-gram model '
        'that exceeds the minimum requirement.',
        'The Advanced task is partially implemented — the GradientBoosting model exists in code but is '
        'not activated in the simulation, and token replay is not used.',
        'Compared to other teams, the process model awareness and Laplace smoothing are unique strengths. '
        'The k-gram model (k=2) provides more context than a bigram model (k=1) while being simpler than '
        'full ML approaches.',
        'The evaluation framework and testing are comprehensive, but activity mapping issues undermine '
        'the comparison results for some gateways.',
        'With the two high-priority improvements (activate advanced mode + add token replay), the '
        'implementation would fully satisfy both basic and advanced requirements.',
    ]
    for c in conclusions:
        doc.add_paragraph(c, style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Final Verdict: ')
    run.bold = True
    run.font.size = Pt(12)
    p.add_run(
        'Good implementation with strong basic task and partial advanced task. '
        'The code quality, testing, and evaluation are above average for an academic project. '
        'Two targeted improvements (token replay + activate advanced mode) would make it complete.'
    )

    # Save
    output_path = os.path.join(os.path.dirname(__file__), 'Task_1_4_Report.docx')
    doc.save(output_path)
    print(f"Report saved to: {output_path}")
    return output_path


if __name__ == '__main__':
    generate_report()
